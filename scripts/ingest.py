from pathlib import Path
import re
import pandas as pd

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import TextLoader, PyPDFLoader
from docx import Document as DocxDocument
import camelot

# ==============================
# PATH CONFIG
# ==============================
RAW_DIR = Path("data/raw")
VECTOR_DIR = Path("faiss_index")

# ==============================
# CLEANING
# ==============================
def clean_text(text: str) -> str:
    text = text.lower()
    text = re.sub(r"\s+", " ", text)
    return text.strip()

# ==============================
# TXT
# ==============================
def load_text(file_path: Path):
    loader = TextLoader(str(file_path), encoding="utf-8")
    return loader.load()

# ==============================
# PDF: TEXT + TABLES
# ==============================
def load_pdf(file_path: Path):
    documents = []

    # 1. Text (no images)
    text_loader = PyPDFLoader(str(file_path))
    documents.extend(text_loader.load())

    # 2. Tables only
    try:
        tables = camelot.read_pdf(str(file_path), pages="all")
        for i, table in enumerate(tables):
            table_text = table.df.to_string(index=False)
            documents.append(
                Document(
                    page_content=table_text,
                    metadata={
                        "source": file_path.name,
                        "type": "pdf_table",
                        "table_index": i
                    }
                )
            )
    except Exception:
        pass  # safe fail if no tables

    return documents

# ==============================
# DOCX: TEXT + TABLES
# ==============================
def load_docx(file_path: Path):
    documents = []
    doc = DocxDocument(file_path)

    # 1. Paragraph text
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if text:
        documents.append(
            Document(
                page_content=text,
                metadata={"source": file_path.name, "type": "docx_text"}
            )
        )

    # 2. Tables
    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        df = pd.DataFrame(rows)
        table_text = df.to_string(index=False)

        documents.append(
            Document(
                page_content=table_text,
                metadata={
                    "source": file_path.name,
                    "type": "docx_table",
                    "table_index": i
                }
            )
        )

    return documents

# ==============================
# CSV (UNCHANGED)
# ==============================
def load_csv(file_path: Path):
    df = pd.read_csv(file_path)
    documents = []
    for _, row in df.iterrows():
        row_text = ", ".join(f"{k}: {v}" for k, v in row.items())

        documents.append(
            Document(
                page_content=row_text.lower(),
                metadata={
                    "source": file_path.name,
                    "type": "csv"
                }
            )
        )

    return documents

# ==============================
# UNIFIED INGEST
# ==============================
def load_documents():
    documents = []

    for file in RAW_DIR.iterdir():
        if file.suffix == ".txt":
            documents.extend(load_text(file))
        elif file.suffix == ".pdf":
            documents.extend(load_pdf(file))
        elif file.suffix == ".docx":
            documents.extend(load_docx(file))
        elif file.suffix == ".csv":
            documents.extend(load_csv(file))

    return documents

# ==============================
# PIPELINE
# ==============================
documents = load_documents()

for doc in documents:
    doc.page_content = clean_text(doc.page_content)

print(f"Loaded {len(documents)} documents")

splitter = RecursiveCharacterTextSplitter(
    chunk_size=700,
    chunk_overlap=150
)

chunks = splitter.split_documents(documents)
print(f"Created {len(chunks)} chunks")

embeddings = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-MiniLM-L6-v2"
)

vectorstore = FAISS.from_documents(chunks, embeddings)
vectorstore.save_local(VECTOR_DIR)

print("FAISS index created and saved")
